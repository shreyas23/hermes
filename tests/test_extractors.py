import socket
from unittest.mock import patch

from bs4 import BeautifulSoup

from extractors import (
    _extract_docx,
    _extract_md,
    _is_safe_url,
    clean_html_for_reader,
    inject_sentence_spans,
    map_images_to_sentences,
)

# --- _is_safe_url (SSRF guard) ---


def test_safe_url_allows_https():
    with patch("socket.gethostbyname", return_value="93.184.216.34"):
        assert _is_safe_url("https://example.com") is True


def test_safe_url_allows_http():
    with patch("socket.gethostbyname", return_value="93.184.216.34"):
        assert _is_safe_url("http://example.com") is True


def test_safe_url_blocks_file_scheme():
    assert _is_safe_url("file:///etc/passwd") is False


def test_safe_url_blocks_ftp_scheme():
    assert _is_safe_url("ftp://example.com") is False


def test_safe_url_blocks_localhost():
    with patch("socket.gethostbyname", return_value="127.0.0.1"):
        assert _is_safe_url("http://localhost/secret") is False


def test_safe_url_blocks_loopback_ip():
    with patch("socket.gethostbyname", return_value="127.0.0.1"):
        assert _is_safe_url("http://127.0.0.1/secret") is False


def test_safe_url_blocks_private_10_range():
    with patch("socket.gethostbyname", return_value="10.0.0.1"):
        assert _is_safe_url("http://10.0.0.1/internal") is False


def test_safe_url_blocks_private_192_range():
    with patch("socket.gethostbyname", return_value="192.168.1.1"):
        assert _is_safe_url("http://192.168.1.1/internal") is False


def test_safe_url_blocks_link_local_metadata():
    with patch("socket.gethostbyname", return_value="169.254.169.254"):
        assert _is_safe_url("http://169.254.169.254/latest/meta-data/") is False


def test_safe_url_blocks_on_dns_failure():
    with patch("socket.gethostbyname", side_effect=socket.gaierror):
        assert _is_safe_url("https://does-not-resolve.invalid") is False


# --- inject_sentence_spans ---


def test_inject_simple_two_sentences():
    html = "<p>First sentence. Second sentence.</p>"
    sentences = ["First sentence.", "Second sentence."]
    out = inject_sentence_spans(html, sentences)
    soup = BeautifulSoup(out, "html.parser")
    spans = soup.find_all("span", attrs={"data-si": True})
    indices = {s["data-si"] for s in spans}
    assert "0" in indices
    assert "1" in indices
    si0 = soup.find("span", attrs={"data-si": "0"})
    assert "First sentence" in si0.get_text()


def test_inject_multiple_sentences_one_text_node():
    html = "<p>Alpha goes here. Beta goes here. Gamma goes here.</p>"
    sentences = ["Alpha goes here.", "Beta goes here.", "Gamma goes here."]
    out = inject_sentence_spans(html, sentences)
    soup = BeautifulSoup(out, "html.parser")
    indices = {s["data-si"] for s in soup.find_all("span", attrs={"data-si": True})}
    assert {"0", "1", "2"}.issubset(indices)


def test_inject_sentence_crossing_inline_element():
    html = "<p>Start of <b>bold sentence.</b> Next one here.</p>"
    sentences = ["Start of bold sentence.", "Next one here."]
    out = inject_sentence_spans(html, sentences)
    soup = BeautifulSoup(out, "html.parser")
    si0 = soup.find("span", attrs={"data-si": "0"})
    assert si0 is not None
    text = si0.get_text()
    assert "Start of" in text
    assert "bold sentence" in text
    # The <b> tag should still be present inside the span
    assert si0.find("b") is not None


def test_inject_unmatched_sentences_skipped_gracefully():
    html = "<p>Only this sentence exists.</p>"
    sentences = ["Only this sentence exists.", "This sentence is not in the html at all."]
    out = inject_sentence_spans(html, sentences)
    soup = BeautifulSoup(out, "html.parser")
    indices = {s["data-si"] for s in soup.find_all("span", attrs={"data-si": True})}
    assert "0" in indices
    assert "1" not in indices


# --- clean_html_for_reader ---


def test_clean_html_strips_script_and_style():
    html = (
        "<html><head><title>My Article</title>"
        "<style>.x{color:red}</style></head>"
        "<body><article><h1>My Article</h1>"
        "<p>This is the body of the article with enough text to be extracted as content here.</p>"
        "<p>A second paragraph so readability keeps the main content block intact and clean.</p>"
        "<script>alert('xss')</script>"
        "</article></body></html>"
    )
    reader_html, _title = clean_html_for_reader(html, "https://example.com")
    assert "<script" not in reader_html.lower()
    assert "<style" not in reader_html.lower()
    assert "alert('xss')" not in reader_html


def test_clean_html_resolves_relative_image_urls():
    html = (
        "<html><head><title>Photo Story</title></head>"
        "<body><article><h1>Photo Story</h1>"
        "<p>Here is a paragraph long enough to count as real article content for readability.</p>"
        "<p>A second paragraph to keep the article body substantial and reliably extracted.</p>"
        '<img src="/img/photo.jpg" width="600" height="400">'
        "</article></body></html>"
    )
    reader_html, _title = clean_html_for_reader(html, "https://example.com")
    soup = BeautifulSoup(reader_html, "html.parser")
    srcs = [img.get("src", "") for img in soup.find_all("img")]
    assert "https://example.com/img/photo.jpg" in srcs


# --- map_images_to_sentences ---


def test_map_images_to_sentences_by_char_ratio():
    sentences = ["First.", "Second.", "Third."]
    images = [
        {"char_offset": 0, "total_chars": 90, "filename": "a.png", "alt": "A"},
        {"char_offset": 60, "total_chars": 90, "filename": "b.png", "alt": "B"},
    ]
    mapped = map_images_to_sentences(images, "", sentences)
    assert len(mapped) == 2
    # ratio 0/90 -> sentence 0
    assert mapped[0]["after_sentence"] == 0
    assert mapped[0]["filename"] == "a.png"
    # ratio 60/90 = 0.66 -> int(0.66 * 3) = 2
    assert mapped[1]["after_sentence"] == 2
    assert mapped[1]["filename"] == "b.png"


# --- _extract_md (structured HTML + TOC) ---

MD_SAMPLE = """# Introduction

Intro paragraph with **bold** and `inline`.

## Details

- one
- two

```python
print("hi")
```

| A | B |
|---|---|
| 1 | 2 |
"""


def _write(tmp_path, name, content, mode="w"):
    p = tmp_path / name
    if mode == "wb":
        p.write_bytes(content)
    else:
        p.write_text(content)
    return str(p)


def test_md_reader_html_has_structure(tmp_path):
    result = _extract_md(_write(tmp_path, "doc.md", MD_SAMPLE))
    html = result["reader_html"]
    assert "<h1" in html
    assert "<h2" in html
    assert "<table" in html
    assert "<code" in html
    assert "<li>" in html


def test_md_toc_is_valid(tmp_path):
    result = _extract_md(_write(tmp_path, "doc.md", MD_SAMPLE))
    toc = result["toc"]
    assert toc == [
        {"level": 1, "title": "Introduction", "id": "h-1"},
        {"level": 2, "title": "Details", "id": "h-2"},
    ]
    # every TOC id resolves to a heading in the HTML
    soup = BeautifulSoup(result["reader_html"], "html.parser")
    for entry in toc:
        assert soup.find(id=entry["id"]) is not None


def test_md_plain_text_is_clean(tmp_path):
    result = _extract_md(_write(tmp_path, "doc.md", MD_SAMPLE))
    text = result["text"]
    assert "<" not in text and ">" not in text
    assert "Introduction" in text
    assert "bold" in text
    # markdown markup chars are gone
    assert "**" not in text and "#" not in text


# --- _extract_docx (structured HTML + TOC) ---


def _make_docx(tmp_path):
    import docx

    d = docx.Document()
    d.add_heading("Doc Title", level=1)
    d.add_paragraph("Body paragraph.")
    d.add_heading("Subsection", level=2)
    d.add_paragraph("Another paragraph.")
    path = str(tmp_path / "doc.docx")
    d.save(path)
    return path


def test_docx_reader_html_has_headings(tmp_path):
    result = _extract_docx(_make_docx(tmp_path))
    html = result["reader_html"]
    assert "<h1" in html
    assert "Doc Title" in html
    assert "<p>" in html


def test_docx_toc_is_valid(tmp_path):
    result = _extract_docx(_make_docx(tmp_path))
    toc = result["toc"]
    titles = [e["title"] for e in toc]
    assert "Doc Title" in titles
    assert "Subsection" in titles
    soup = BeautifulSoup(result["reader_html"], "html.parser")
    for entry in toc:
        assert soup.find(id=entry["id"]) is not None


def test_docx_plain_text_is_clean(tmp_path):
    result = _extract_docx(_make_docx(tmp_path))
    text = result["text"]
    assert "<" not in text and ">" not in text
    assert "Doc Title" in text
    assert "Body paragraph." in text
