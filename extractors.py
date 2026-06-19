import hashlib
import os
import re
from pathlib import Path
from urllib.parse import urljoin

SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.md', '.txt', '.rtf', '.html', '.htm'}


def extract_text(file_path: str) -> str | None:
    result = extract_with_images(file_path)
    return result['text'] if result else None


def extract_with_images(file_path: str, image_dir: str = None) -> dict | None:
    ext = Path(file_path).suffix.lower()
    extractors = {
        '.txt': _extract_txt,
        '.md': _extract_md,
        '.html': _extract_html,
        '.htm': _extract_html,
        '.pdf': _extract_pdf,
        '.docx': _extract_docx,
        '.rtf': _extract_rtf,
    }
    fn = extractors.get(ext)
    if not fn:
        return None
    try:
        return fn(file_path, image_dir)
    except Exception as e:
        print(f"Extraction error for {file_path}: {e}")
        return None


def clean_html_for_reader(html: str, base_url: str) -> str:
    from readability import Document
    from bs4 import BeautifulSoup

    # Pre-process: unwrap <picture> to expose <img>, capture dimensions
    orig_soup = BeautifulSoup(html, 'html.parser')
    img_dims = {}
    for img in orig_soup.find_all('img'):
        src = img.get('src', '')
        if src:
            dims = {}
            if img.get('height') and img['height']:
                dims['height'] = str(img['height'])
            if img.get('width') and img['width']:
                dims['width'] = str(img['width'])
            if dims:
                img_dims[src] = dims

    for picture in orig_soup.find_all('picture'):
        img = picture.find('img')
        if img:
            picture.replace_with(img)
        else:
            source = picture.find('source')
            if source and source.get('srcset'):
                new_img = orig_soup.new_tag('img', src=source['srcset'].split()[0])
                picture.replace_with(new_img)

    doc = Document(str(orig_soup))
    reader_html = doc.summary()
    doc_title = doc.title()
    if doc_title == '[no-title]':
        doc_title = None

    soup = BeautifulSoup(reader_html, 'html.parser')

    # Re-add title that readability extracts separately
    if doc_title:
        content = soup.find('div') or soup
        first_child = content.find()
        if first_child:
            h1 = soup.new_tag('h1')
            h1.string = doc_title
            first_child.insert_before(h1)

    for tag in soup.find_all(True):
        tag.attrs = {k: v for k, v in tag.attrs.items() if v}
    for img in soup.find_all('img'):
        src = img.get('src', '')
        if src:
            abs_src = urljoin(base_url, src)
            img['src'] = abs_src
            # Restore original dimensions
            orig = img_dims.get(src) or img_dims.get(abs_src) or {}
            for attr in ('height', 'width'):
                if attr in orig and attr not in img.attrs:
                    img[attr] = orig[attr]
    for a in soup.find_all('a'):
        href = a.get('href', '')
        if href:
            a['href'] = urljoin(base_url, href)

    # Re-inject content images that readability dropped
    def _img_key(url):
        """Extract base image path, ignoring CDN resize params."""
        import re
        m = re.search(r'([\w-]{20,}\.\w{3,4})$', url.split('?')[0].split('%2F')[-1])
        return m.group(1) if m else url

    reader_img_keys = {_img_key(img.get('src', '')) for img in soup.find_all('img')}
    article = orig_soup.find('article') or orig_soup.find('main') or orig_soup
    for img in article.find_all('img'):
        src = img.get('src', '')
        abs_src = urljoin(base_url, src) if src else ''
        key = _img_key(abs_src)
        if not abs_src or key in reader_img_keys:
            continue
        w = str(img.get('width', '999')).split('.')[0]
        h = str(img.get('height', '999')).split('.')[0]
        if (w.isdigit() and int(w) < 100) or (h.isdigit() and int(h) < 50):
            continue

        alt = img.get('alt', '')
        prev_text = ''
        for sib in img.previous_siblings:
            t = getattr(sib, 'get_text', lambda: str(sib))()
            if t.strip():
                prev_text = t.strip()
                break

        new_img = soup.new_tag('img', src=abs_src)
        if alt:
            new_img['alt'] = alt
        if w.isdigit() and int(w) > 0:
            new_img['width'] = w
        if h.isdigit() and int(h) > 0:
            new_img['height'] = h

        inserted = False
        if prev_text:
            words = prev_text[-40:].split()
            search = ' '.join(words[-4:]) if len(words) >= 4 else prev_text[-20:]
            for tn in soup.find_all(string=True):
                if search in str(tn):
                    parent = tn.find_parent()
                    if parent:
                        parent.insert_after(new_img)
                        inserted = True
                        break

        if not inserted:
            paras = soup.find_all('p')
            if paras:
                ratio = len(reader_img_keys) / max(len(paras), 1)
                insert_at = min(int(ratio * len(paras)), len(paras) - 1)
                paras[insert_at].insert_after(new_img)

        reader_img_keys.add(key)

    return str(soup)


def extract_url_with_images(html: str, base_url: str, image_dir: str = None) -> dict | None:
    from bs4 import BeautifulSoup
    import urllib.request

    soup = BeautifulSoup(html, 'html.parser')
    for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
        tag.decompose()

    body = soup.find('article') or soup.find('main') or soup.find('body') or soup
    blocks = []
    char_offset = 0

    for el in body.descendants:
        if el.name == 'img':
            src = el.get('src', '')
            if not src or src.startswith('data:'):
                continue
            src = urljoin(base_url, src)
            alt = el.get('alt', '')
            img_entry = {'type': 'image', 'src': src, 'alt': alt, 'char_offset': char_offset}
            if image_dir:
                local = _download_image(src, image_dir, len([b for b in blocks if b['type'] == 'image']))
                if local:
                    img_entry['filename'] = local
            blocks.append(img_entry)
        elif el.string and el.parent.name not in ('script', 'style', 'img'):
            text = el.string.strip()
            if text:
                blocks.append({'type': 'text', 'content': text, 'char_offset': char_offset})
                char_offset += len(text) + 1

    full_text = ' '.join(b['content'] for b in blocks if b['type'] == 'text')
    total_chars = char_offset or 1
    images = []
    for b in blocks:
        if b['type'] == 'image' and b.get('filename'):
            b['total_chars'] = total_chars
            images.append(b)

    return {'text': full_text, 'images': images}


def map_images_to_sentences(images: list, text: str, sentences: list) -> list:
    if not images or not sentences:
        return []

    mapped = []
    for img in images:
        char_pos = img.get('char_offset', 0)
        total_chars = img.get('total_chars', 1) or 1
        ratio = char_pos / total_chars
        after_sentence = min(int(ratio * len(sentences)), len(sentences) - 1)
        mapped.append({
            'after_sentence': after_sentence,
            'filename': img.get('filename', ''),
            'alt': img.get('alt', ''),
        })

    return mapped


def _download_image(url: str, image_dir: str, index: int) -> str | None:
    import urllib.request
    try:
        ext = _guess_image_ext(url)
        filename = f'img_{index}{ext}'
        filepath = os.path.join(image_dir, filename)
        urllib.request.urlretrieve(url, filepath)
        if os.path.getsize(filepath) < 100:
            os.unlink(filepath)
            return None
        return filename
    except Exception as e:
        print(f"Failed to download image {url}: {e}")
        return None


def _guess_image_ext(url: str) -> str:
    url_lower = url.split('?')[0].lower()
    for ext in ('.png', '.jpg', '.jpeg', '.gif', '.webp', '.svg'):
        if url_lower.endswith(ext):
            return ext
    return '.jpg'


# --- Format extractors (return dict with text + images) ---

def _extract_txt(path: str, image_dir: str = None) -> dict:
    with open(path, 'r', errors='ignore') as f:
        return {'text': f.read(), 'images': []}


def _extract_md(path: str, image_dir: str = None) -> dict:
    with open(path, 'r', errors='ignore') as f:
        text = f.read()

    images = []
    if image_dir:
        for i, match in enumerate(re.finditer(r'!\[([^\]]*)\]\(([^)]+)\)', text)):
            alt, src = match.group(1), match.group(2)
            if src.startswith('http'):
                local = _download_image(src, image_dir, i)
            elif os.path.isfile(os.path.join(os.path.dirname(path), src)):
                import shutil
                ext = Path(src).suffix
                local = f'img_{i}{ext}'
                shutil.copy2(os.path.join(os.path.dirname(path), src), os.path.join(image_dir, local))
            else:
                continue
            if local:
                images.append({'type': 'image', 'filename': local, 'alt': alt, 'char_offset': match.start(), 'total_chars': len(text) or 1})

    text = re.sub(r'```[\s\S]*?```', '', text)
    text = re.sub(r'`[^`]+`', '', text)
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    text = re.sub(r'\[([^\]]+)\]\(.*?\)', r'\1', text)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'[*_]{1,3}([^*_]+)[*_]{1,3}', r'\1', text)
    text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*>\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return {'text': text.strip(), 'images': images}


def _extract_html(path: str, image_dir: str = None) -> dict:
    with open(path, 'r', errors='ignore') as f:
        html = f.read()
    return extract_url_with_images(html, f'file://{path}', image_dir)


def _extract_pdf(path: str, image_dir: str = None) -> dict:
    import pymupdf
    doc = pymupdf.open(path)
    pages = []
    images = []
    char_offset = 0

    for page_num, page in enumerate(doc):
        text = page.get_text()
        if text and text.strip():
            pages.append(text.strip())

        if image_dir:
            for i, img in enumerate(page.get_images(full=True)):
                try:
                    xref = img[0]
                    pix = pymupdf.Pixmap(doc, xref)
                    if pix.n > 4:
                        pix = pymupdf.Pixmap(pymupdf.csRGB, pix)
                    filename = f'img_p{page_num}_{i}.png'
                    pix.save(os.path.join(image_dir, filename))
                    if os.path.getsize(os.path.join(image_dir, filename)) > 500:
                        images.append({
                            'type': 'image',
                            'filename': filename,
                            'alt': f'Page {page_num + 1} image',
                            'char_offset': char_offset + len(text or '') // 2,
                        })
                except Exception:
                    pass

        char_offset += len(text or '') + 2

    doc.close()
    total_chars = char_offset or 1
    for img in images:
        img['total_chars'] = total_chars
    return {'text': '\n\n'.join(pages), 'images': images}


def _extract_docx(path: str, image_dir: str = None) -> dict:
    import docx
    doc = docx.Document(path)
    texts = []
    images = []
    char_offset = 0

    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

        if image_dir:
            for run in para.runs:
                for rel in run.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    blips = rel.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    for blip in blips:
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed and embed in doc.part.rels:
                            image_part = doc.part.rels[embed].target_part
                            ext = Path(image_part.partname).suffix
                            filename = f'img_{len(images)}{ext}'
                            with open(os.path.join(image_dir, filename), 'wb') as f:
                                f.write(image_part.blob)
                            images.append({
                                'type': 'image',
                                'filename': filename,
                                'alt': '',
                                'char_offset': char_offset,
                            })

        char_offset += len(para.text) + 2

    total_chars = char_offset or 1
    for img in images:
        img['total_chars'] = total_chars
    return {'text': '\n\n'.join(texts), 'images': images}


def _extract_rtf(path: str, image_dir: str = None) -> dict:
    from striprtf.striprtf import rtf_to_text
    with open(path, 'r', errors='ignore') as f:
        return {'text': rtf_to_text(f.read()), 'images': []}
